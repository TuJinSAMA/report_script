import datetime
import os
import glob
import re
import time

import win32com.client as win32
from win32com.client import constants
from docx import Document
from lxml import etree

report_dir = r'E:\Temp\0825-fix'


# report_dir = r'E:\Temp\188人word'
# output_dir = r'E:\Temp\0825-fix'


def fix_docx(docx_path):
    doc = Document(docx_path)
    p_list = []
    for paragraph in doc.paragraphs:
        p_list.append(paragraph.text)
        # if paragraph.text.strip() == 'A':
        #
        #     paragraph.clear()
        #     # 添加换行
        #     for _ in range(6):  # 这里的5可以根据需要修改换行的数量
        #         paragraph.add_run().add_break()
    if "陈其栋" in docx_path or "王轩" in docx_path:
        # print(len(p_list))
        # print('\n')
        # print(p_list)
        # print('\n')
        print(docx_path)
        print('\n')
        for table in doc.tables:
            text = table.rows[0].cells[0].text
            print(table)

    # doc.save(os.path.join(output_dir, os.path.basename(docx_path)))


def fix(report_dir):
    for file_name in os.listdir(report_dir):
        if file_name.endswith('.docx'):
            # 获取docx 文件
            docx_path = os.path.join(report_dir, file_name)
            fix_docx(docx_path)


if __name__ == '__main__':
    # exec_script()
    time1 = time.time()
    fix(report_dir)
    time2 = time.time()
    print(f"{time2 - time1}")

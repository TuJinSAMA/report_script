import os
import glob
import re
import win32com.client as win32
from win32com.client import constants
from docx import Document
from lxml import etree

# report_dir = r'D:\Temp\Report\个体报告批量导出'
# template_file = r'D:\Temp\Report\template.docx'
# output_dir = r'D:\Temp\Report\output'

# report_dir = r'E:\Temp\个体报告批量导出'
report_dir = r'E:\Temp\2025-0706'
template_file = r'E:\Temp\template2.docx'
output_dir = r'E:\Temp\output-2025-0706'

crxl = r'成人心理压力量表'
askrg = r'艾森克人格测验'
jlzp = r'焦虑自评量表'
zzzp = r'症状自评量表'
zwhx = r'自我和谐量表'


def extract_info(text):
    # 使用 split 方法按冒号分割字符串，并取第二部分
    name = text.split("：")[1] if "：" in text else ""
    # 输出提取的姓名
    return name


def dispose_askrg(doc, template):
    # 遍历文档中的所有表格
    for table in doc.tables:
        # 检查表格的第一行是否包含目标表头
        if '维度' in table.rows[0].cells[0].text:
            # 找到目标表格，可以进行操作
            # 定义正则表达式模式，匹配 "分" 后插入冒号
            pattern = r"分(?!：)"  # 这确保我们不会重复插入冒号
            # 使用 re.sub 进行替换，只替换第一个匹配
            raw = re.sub(pattern, "分：", table.rows[1].cells[0].text, count=1)
            standard = re.sub(pattern, "分：", table.rows[1].cells[1].text, count=1)
            result = table.rows[2].cells[0].text
            suggest = table.rows[3].cells[0].text
            # 操作：修改表格的内容
            for t in template.tables:
                print(t.rows[0].cells[0].text)
                if askrg in t.rows[0].cells[0].text:
                    if '维度一' in table.rows[0].cells[0].text:
                        t.rows[2].cells[0].text = raw
                        t.rows[2].cells[1].text = standard
                        t.rows[3].cells[0].text = result
                        t.rows[4].cells[0].text = suggest
                        print(t.rows[2].cells[0].text)
                        print(t.rows[2].cells[1].text)
                        print(t.rows[3].cells[0].text)
                        print(t.rows[4].cells[0].text)
                    if '维度二' in table.rows[0].cells[0].text:
                        t.rows[6].cells[0].text = raw
                        t.rows[6].cells[1].text = standard
                        t.rows[7].cells[0].text = result
                        t.rows[8].cells[0].text = suggest
                        print(t.rows[6].cells[0].text)
                        print(t.rows[6].cells[1].text)
                        print(t.rows[7].cells[0].text)
                        print(t.rows[8].cells[0].text)
                    if '维度三' in table.rows[0].cells[0].text:
                        t.rows[10].cells[0].text = raw
                        t.rows[10].cells[1].text = standard
                        t.rows[11].cells[0].text = result
                        t.rows[12].cells[0].text = suggest
                        print(t.rows[10].cells[0].text)
                        print(t.rows[10].cells[1].text)
                        print(t.rows[11].cells[0].text)
                        print(t.rows[12].cells[0].text)
                    if '维度四' in table.rows[0].cells[0].text:
                        t.rows[14].cells[0].text = raw
                        t.rows[14].cells[1].text = standard
                        t.rows[15].cells[0].text = result
                        t.rows[16].cells[0].text = suggest
                        print(t.rows[14].cells[0].text)
                        print(t.rows[14].cells[1].text)
                        print(t.rows[15].cells[0].text)
                        print(t.rows[16].cells[0].text)
    return


def dispose_crxl(doc, template):
    name = ''
    gender = ''
    birthday = ''
    completed_date = ''
    # 遍历文档中的所有表格
    for table in doc.tables:
        if '登录名' in table.rows[0].cells[0].text:
            name = extract_info(table.rows[0].cells[1].text)
            gender = extract_info(table.rows[0].cells[2].text)
            birthday = extract_info(table.rows[1].cells[0].text)
            completed_date = extract_info(table.rows[1].cells[1].text)
        # 检查表格的第一行是否包含目标表头
        if '总评' in table.rows[0].cells[0].text:
            # 找到目标表格，可以进行操作
            # 定义正则表达式模式，匹配 "分" 后插入冒号
            pattern = r"分(?!：)"  # 这确保我们不会重复插入冒号
            # 使用 re.sub 进行替换，只替换第一个匹配
            raw = re.sub(pattern, "分：", table.rows[1].cells[0].text, count=1)
            standard = re.sub(pattern, "分：", table.rows[1].cells[1].text, count=1)
            result = table.rows[2].cells[0].text
            suggest = table.rows[3].cells[0].text
            # 操作：修改表格的内容
            for t in template.tables:
                if '姓名' in t.rows[0].cells[0].text:
                    cell_list = [
                        {
                            'cell': t.rows[0].cells[1],
                            'value': name
                        },
                        {
                            'cell': t.rows[0].cells[3],
                            'value': gender
                        },
                        {
                            'cell': t.rows[1].cells[1],
                            'value': birthday
                        },
                        {
                            'cell': t.rows[1].cells[3],
                            'value': completed_date
                        },
                    ]
                    for item in cell_list:
                        # 检查单元格中是否有段落
                        if len(item['cell'].paragraphs) > 0:
                            # 获取第一个段落
                            paragraph = item['cell'].paragraphs[0]
                            # 清空段落中的所有文本块
                            for run in paragraph.runs:
                                run.text = ''
                            # 添加新的文本块，并设置文本内容
                            paragraph.add_run(item['value'])

                    # t.rows[0].cells[3].text = gender
                    # t.rows[1].cells[1].text = birthday
                if crxl in t.rows[0].cells[0].text:
                    t.rows[1].cells[0].text = raw
                    t.rows[1].cells[1].text = standard
                    t.rows[2].cells[0].text = result
                    t.rows[3].cells[0].text = suggest
                    print(t.rows[1].cells[0].text)
                    print(t.rows[1].cells[1].text)
                    print(t.rows[2].cells[0].text)
                    print(t.rows[3].cells[0].text)
    return name, birthday


def dispose_jlzp(doc, template):
    # 遍历文档中的所有表格
    for table in doc.tables:
        # 检查表格的第一行是否包含目标表头
        if '总评' in table.rows[0].cells[0].text:
            # 找到目标表格，可以进行操作
            # 定义正则表达式模式，匹配 "分" 后插入冒号
            pattern = r"分(?!：)"  # 这确保我们不会重复插入冒号
            # 使用 re.sub 进行替换，只替换第一个匹配
            raw = re.sub(pattern, "分：", table.rows[1].cells[0].text, count=1)
            standard = re.sub(pattern, "分：", table.rows[1].cells[1].text, count=1)
            result = table.rows[2].cells[0].text
            suggest = table.rows[3].cells[0].text
            # 操作：修改表格的内容
            for t in template.tables:
                print(t.rows[0].cells[0].text)
                if jlzp in t.rows[0].cells[0].text:
                    t.rows[1].cells[0].text = raw
                    t.rows[1].cells[1].text = standard
                    t.rows[2].cells[0].text = result
                    t.rows[3].cells[0].text = suggest
                    print(t.rows[1].cells[0].text)
                    print(t.rows[1].cells[1].text)
                    print(t.rows[2].cells[0].text)
                    print(t.rows[3].cells[0].text)
                    break  # 找到目标表格后退出循环
            break
    return


def dispose_zzzp(doc, template):
    # 遍历文档中的所有表格
    for table in doc.tables:
        # 检查表格的第一行是否包含目标表头
        print(table.rows[0].cells[0].text)
        if '维度一' in table.rows[0].cells[0].text:
            # 找到目标表格，可以进行操作
            # 定义正则表达式模式，匹配 "分" 后插入冒号
            pattern = r"分(?!：)"  # 这确保我们不会重复插入冒号
            # 使用 re.sub 进行替换，只替换第一个匹配
            raw = re.sub(pattern, "分：", table.rows[1].cells[0].text, count=1)
            standard = re.sub(pattern, "分：", table.rows[1].cells[1].text, count=1)
            result = table.rows[2].cells[0].text
            result = result.split('\n')
            result = result[0] + '\n' + result[2]
            # result = table.rows[2].cells[0].text
            suggest = table.rows[3].cells[0].text
            # 操作：修改表格的内容
            for t in template.tables:
                print(t.rows[0].cells[0].text)
                if zzzp in t.rows[0].cells[0].text:
                    t.rows[1].cells[0].text = raw
                    t.rows[1].cells[1].text = standard
                    t.rows[2].cells[0].text = result
                    t.rows[3].cells[0].text = suggest
                    print(t.rows[1].cells[0].text)
                    print(t.rows[1].cells[1].text)
                    print(t.rows[2].cells[0].text)
                    print(t.rows[3].cells[0].text)
                    break  # 找到目标表格后退出循环
            break
    return


def dispose_zwhx(doc, template):
    # 遍历文档中的所有表格
    for table in doc.tables:
        # 检查表格的第一行是否包含目标表头
        print(table.rows[0].cells[0].text)
        if '总评' in table.rows[0].cells[0].text:
            # 找到目标表格，可以进行操作
            # 定义正则表达式模式，匹配 "分" 后插入冒号
            pattern = r"分(?!：)"  # 这确保我们不会重复插入冒号
            # 使用 re.sub 进行替换，只替换第一个匹配
            raw = re.sub(pattern, "分：", table.rows[1].cells[0].text, count=1)
            standard = re.sub(pattern, "分：", table.rows[1].cells[1].text, count=1)
            result = table.rows[2].cells[0].text
            suggest = table.rows[3].cells[0].text
            # 操作：修改表格的内容
            for t in template.tables:
                print(t.rows[0].cells[0].text)
                if zwhx in t.rows[0].cells[0].text:
                    t.rows[1].cells[0].text = raw
                    t.rows[1].cells[1].text = standard
                    t.rows[2].cells[0].text = result
                    t.rows[3].cells[0].text = suggest
                    print(t.rows[1].cells[0].text)
                    print(t.rows[1].cells[1].text)
                    print(t.rows[2].cells[0].text)
                    print(t.rows[3].cells[0].text)
                    break  # 找到目标表格后退出循环
            break
    return


def read_docx(path, template_doc):
    # 打开 Word 文档
    doc = Document(path)

    file_name = ''
    if re.search(crxl, path):
        print(f"开始处理: {path}")
        name, birthday = dispose_crxl(doc, template_doc)
        file_name = f'{name}-{birthday}.docx'
    if re.search(askrg, path):
        print(f"开始处理: {path}")
        dispose_askrg(doc, template_doc)
    if re.search(jlzp, path):
        print(f"开始处理: {path}")
        dispose_jlzp(doc, template_doc)
    if re.search(zzzp, path):
        print(f"开始处理: {path}")
        dispose_zzzp(doc, template_doc)
    if re.search(zwhx, path):
        print(f"开始处理: {path}")
        dispose_zwhx(doc, template_doc)
    return file_name


def change_textbox(file_path):
    doc_app = win32.gencache.EnsureDispatch('Word.Application')  # 打开word应用程序
    doc_app.Visible = False  # 设置应用程序可见
    doc = doc_app.Documents.Open(file_path)
    # 获取文件名部分
    file_name = os.path.basename(file_path)
    # 查找第一个"-"字符的位置
    dash_index = file_name.find("-")
    name = ''
    if dash_index != -1:
        # 提取"-"之前的文字作为人名
        name = file_name[:dash_index]
        print("提取的人名:", name)
    else:
        print("文件名中没有找到'-'字符")
    for shape in doc.Shapes:
        if shape.Type == 17 and shape.TextFrame.HasText:
            if '{{name}}' in shape.TextFrame.TextRange.Text:
                shape.TextFrame.TextRange.Text = shape.TextFrame.TextRange.Text.replace('{{name}}', name)
    doc.Save()
    doc.Close()
    doc_app.Quit()


def exec_script(path):
    output_path = ''
    docx_files = glob.glob(os.path.join(path, "*.docx"))
    template_doc = Document(template_file)
    for file in docx_files:
        print(file)
        file_name = read_docx(file, template_doc)
        if file_name != '':
            output_path = os.path.join(output_dir, file_name)
    template_doc.save(output_path)
    change_textbox(output_path)
    return


if __name__ == '__main__':
    # exec_script()
    for dir_name in os.listdir(report_dir):
        # 获取完整的文件夹路径
        dir_path = os.path.join(report_dir, dir_name)
        # 检查是否为文件夹
        if os.path.isdir(dir_path):
            print(f"找到文件夹: {dir_name}")
            exec_script(dir_path)
